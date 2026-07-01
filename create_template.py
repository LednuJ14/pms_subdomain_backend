from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os

doc = Document()

# Define global styles to make it professional and strictly black
style_normal = doc.styles['Normal']
font_normal = style_normal.font
font_normal.name = 'Arial'
font_normal.size = Pt(11)
font_normal.color.rgb = RGBColor(0, 0, 0)

title_style = doc.styles['Title']
title_style.font.name = 'Arial'
title_style.font.size = Pt(18)
title_style.font.bold = True
title_style.font.color.rgb = RGBColor(0, 0, 0)
title_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
title_style.paragraph_format.space_after = Pt(24)

h1_style = doc.styles['Heading 1']
h1_style.font.name = 'Arial'
h1_style.font.size = Pt(14)
h1_style.font.bold = True
h1_style.font.color.rgb = RGBColor(0, 0, 0)
h1_style.paragraph_format.space_before = Pt(18)
h1_style.paragraph_format.space_after = Pt(6)

# Title
doc.add_paragraph('RENTAL CONTRACT AGREEMENT', style='Title')

# Meta info block
p_meta = doc.add_paragraph()
p_meta.add_run('Contract Number: ').bold = True
p_meta.add_run('{{ contract_number }}\n')
p_meta.add_run('Date Generated: ').bold = True
p_meta.add_run('{{ date_generated }}')
p_meta.paragraph_format.space_after = Pt(24)

# 1. Parties
doc.add_heading('1. Parties', level=1)
p1 = doc.add_paragraph('This Rental Contract Agreement is made and entered into between:\n\n')
p1.add_run('Landlord / Property Manager: ').bold = True
p1.add_run('{{ property_name }}\n')
p1.add_run('Tenant: ').bold = True
p1.add_run('{{ tenant_name }} ({{ tenant_email }})')
p1.paragraph_format.space_after = Pt(14)

# 2. Property Details
doc.add_heading('2. Property Details', level=1)
doc.add_paragraph('The Landlord agrees to rent to the Tenant the following premises:')
p2 = doc.add_paragraph()
p2.add_run('Property Name: ').bold = True
p2.add_run('{{ property_name }}\n')
p2.add_run('Address: ').bold = True
p2.add_run('{{ property_address }}\n')
p2.add_run('Unit / Room: ').bold = True
p2.add_run('{{ unit_number }}')
p2.paragraph_format.left_indent = Inches(0.5)
p2.paragraph_format.space_after = Pt(14)

# 3. Term of Lease
doc.add_heading('3. Term of Lease', level=1)
doc.add_paragraph('The lease shall commence on {{ start_date }} and terminate on {{ end_date }} (Type: {{ contract_type }}).').paragraph_format.space_after = Pt(14)

# 4. Financial Terms
doc.add_heading('4. Financial Terms', level=1)
p4 = doc.add_paragraph()
p4.add_run('Monthly Rent: ').bold = True
p4.add_run('{{ monthly_rent }}\n')
p4.add_run('Security Deposit: ').bold = True
p4.add_run('{{ security_deposit }}\n')
p4.add_run('Total Value: ').bold = True
p4.add_run('{{ total_contract_value }}')
p4.paragraph_format.left_indent = Inches(0.5)
p4.paragraph_format.space_after = Pt(14)

# 5. Terms and Conditions
doc.add_heading('5. Terms and Conditions', level=1)
doc.add_paragraph('{{ terms_and_conditions }}').paragraph_format.space_after = Pt(14)

# 6. Special Conditions
doc.add_heading('6. Special Conditions', level=1)
doc.add_paragraph('{{ special_conditions }}').paragraph_format.space_after = Pt(14)

# 7. Signatures
doc.add_heading('7. Signatures', level=1)
doc.add_paragraph('By signing below, both parties agree to the terms outlined above.').paragraph_format.space_after = Pt(24)

# Signature block
table = doc.add_table(rows=2, cols=2)
table.autofit = True
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = ''
hdr_cells[0].paragraphs[0].add_run('LANDLORD / MANAGER').bold = True
hdr_cells[1].text = ''
hdr_cells[1].paragraphs[0].add_run('TENANT').bold = True

row_cells = table.rows[1].cells
row_cells[0].text = '{% if landlord_signed %}Signed ({{ landlord_signed_date }})\n{{ landlord_signature_image }}{% else %}Not Signed{% endif %}'
row_cells[1].text = '{% if tenant_signed %}Signed ({{ tenant_signed_date }})\n{{ tenant_signature_image }}{% else %}Not Signed{% endif %}'

# Padding for cells
for row in table.rows:
    for cell in row.cells:
        cell.paragraphs[0].paragraph_format.space_before = Pt(6)
        cell.paragraphs[0].paragraph_format.space_after = Pt(6)

doc.add_page_break()

# Audit Trail
doc.add_heading('E-Signature Audit Trail', level=1)
doc.add_paragraph('This document was electronically signed via the JACS Property Management System. The following information serves as the legal audit trail for this document.').paragraph_format.space_after = Pt(14)

audit_table = doc.add_table(rows=4, cols=3)
audit_table.style = 'Table Grid'
h_cells = audit_table.rows[0].cells
headers = ['Role', 'IP Address', 'User Agent / Device']
for i, header in enumerate(headers):
    h_cells[i].text = ''
    h_cells[i].paragraphs[0].add_run(header).bold = True

r1_cells = audit_table.rows[1].cells
r1_cells[0].text = 'Landlord'
r1_cells[1].text = '{{ landlord_ip }}'
r1_cells[2].text = '{{ landlord_user_agent }}'

r2_cells = audit_table.rows[2].cells
r2_cells[0].text = 'Tenant'
r2_cells[1].text = '{{ tenant_ip }}'
r2_cells[2].text = '{{ tenant_user_agent }}'

r3_cells = audit_table.rows[3].cells
r3_cells[0].text = 'Document Hash'
r3_cells[1].text = '{{ document_hash }}'
r3_cells[2].text = ''

for row in audit_table.rows:
    for cell in row.cells:
        cell.paragraphs[0].paragraph_format.space_before = Pt(6)
        cell.paragraphs[0].paragraph_format.space_after = Pt(6)

os.makedirs('templates', exist_ok=True)
doc.save('templates/default_contract_template.docx')
print("Professional template created successfully!")
